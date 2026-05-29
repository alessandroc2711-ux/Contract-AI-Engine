import os
import re
import json
import streamlit as st
import pdfplumber
import pytesseract
from pdf2image import convert_from_path
from docx import Document
from openai import OpenAI

# =========================
# INIT CLIENT
# =========================
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# =========================
# SESSION STATE
# =========================
def init_state():
    defaults = {
        "template_contract": None,
        "example_contract": None,
        "example_request": None,
        "example_event": None,
        "new_request": None,
        "new_event": None,

        "mapping": None,
        "extracted_fields": None,
        "final_docx": None,       
    }

    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


def require(key, msg):
    if not st.session_state.get(key):
        st.error(msg)
        st.stop()


init_state()

# =========================
# AI ENGINE
# =========================
class AIEngine:

    def safe_json_loads(self, text):
        try:
            return json.loads(text)
        except Exception:
            text = str(text).strip()

            text = re.sub(r"^```json", "", text)
            text = re.sub(r"^```", "", text)
            text = re.sub(r"```$", "", text)

            text = text.strip()

            text = re.sub(r'\\(?!["\\/bfnrt])', r'\\\\', text)

            try:
                return json.loads(text)
            except Exception:
                return {"error": "invalid_json", "raw": text}

    def build_mapping(self, template_fields, example_contract, example_request, example_event):

        prompt = f"""
Return ONLY JSON.

You are building a FIELD MAPPING SYSTEM.

IMPORTANT:
- Every template field MUST be mapped.
- Preserve original field names EXACTLY.
- Do not omit fields even if uncertain.

=== TEMPLATE FIELDS ===
{template_fields}

=== EXAMPLE CONTRACT ===
{example_contract}

=== EXAMPLE REQUEST ===
{example_request}

=== EXAMPLE EVENT ===
{example_event}

OUTPUT:
{{
  "fields": [
    {{
      "field_name": "...",
      "source": "REQUEST|EVENT",
      "semantic_meaning": "...",
      "rule": "optional"
    }}
  ]
}}
"""

        res = client.chat.completions.create(
            model="gpt-5",
            messages=[
                {"role": "system", "content": "Return ONLY valid JSON."},
                {"role": "user", "content": prompt}
            ]
        )

        result = self.safe_json_loads(res.choices[0].message.content)

        st.session_state["debug_mapping_raw_fields"] = result

        return result

    def extract_fields(self, mapping, new_request, new_event):

        allowed_fields = [
            f["field_name"]
            for f in mapping.get("fields", [])
            if f.get("field_name")
        ]

        prompt = f"""
Return ONLY JSON.

IMPORTANT:
- Try to extract ALL known fields.
- Preserve field names EXACTLY.
- Never rename fields.
- Never invent new field names.

REQUEST:
{new_request}

EVENT:
{new_event}

KNOWN FIELDS:
{json.dumps(mapping.get("fields", []), indent=2)}

OUTPUT:
{{
  "fields": [
    {{
      "field_name": "...",
      "value": "...",
      "evidence": "...",
      "source": "REQUEST|EVENT"
    }}
  ]
}}
"""

        res = client.chat.completions.create(
            model="gpt-5",
            messages=[
                {"role": "system", "content": "Return ONLY valid JSON."},
                {"role": "user", "content": prompt}
            ]
        )

        raw = res.choices[0].message.content

        st.session_state["debug_extraction_raw_output"] = raw

        result = self.safe_json_loads(raw)

        if not isinstance(result, dict):
            result = {"fields": []}

        result.setdefault("fields", [])

        st.session_state["debug_extraction_before_filter"] = result

        filtered = [
            f for f in result["fields"]
            if f.get("field_name") in allowed_fields
        ]

        result["fields"] = filtered

        st.session_state["debug_extraction_after_filter"] = result

        return result


# =========================
# DOCUMENT PARSER
# =========================
class DocumentParser:

    def parse(self, path):

        if not path:
            return {"text": ""}

        ext = os.path.splitext(path)[-1].lower()

        if ext == ".pdf":
            return self._pdf(path)

        if ext == ".docx":
            return self._docx(path)

        return self._txt(path)

    def _pdf(self, path):

        text_layer = self._text_layer(path)
        ocr_text = self._ocr(path)

        return {
            "text": self._merge(text_layer, ocr_text)
        }

    def _text_layer(self, path):

        out = []

        with pdfplumber.open(path) as pdf:

            for p in pdf.pages:

                words = p.extract_words()

                t = (
                    " ".join(w.get("text", "") for w in words)
                    if words else None
                )

                if not t:
                    t = p.extract_text()

                if not t:
                    t = p.extract_text(layout=True)

                if not t:
                    t = p.extract_text(
                        x_tolerance=2,
                        y_tolerance=2
                    )

                if t:
                    out.append(t)

        return "\n".join(out)

    def _ocr(self, path):

        try:
            imgs = convert_from_path(path, dpi=300)
        except Exception:
            return ""

        return "\n".join(
            pytesseract.image_to_string(i)
            for i in imgs
        )

    def _merge(self, a, b):
        return "\n".join(
            dict.fromkeys(
                (a + "\n" + b).splitlines()
            )
        )

    def _docx(self, path):

        doc = Document(path)

        parts = []

        for p in doc.paragraphs:
            parts.append(
                "".join(run.text for run in p.runs)
            )

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        parts.append(
                            "".join(run.text for run in p.runs)
                        )

        return {"text": "\n".join(parts)}

    def _txt(self, path):

        try:
            with open(path, "r", encoding="utf-8") as f:
                return {"text": f.read()}
        except Exception:
            return {"text": ""}


parser = DocumentParser()
ai = AIEngine()

# =========================
# HELPERS
# =========================
def upload(label, key, state):

    f = st.file_uploader(label, key=key)

    if f:

        path = os.path.join("temp", f.name)

        with open(path, "wb") as w:
            w.write(f.read())

        st.session_state[state] = path


def parse_text(path):
    return parser.parse(path)["text"]


def extract_template_fields(text):

    return sorted(
        set(
            re.findall(
                r"\{\{\s*([^}]+?)\s*\}\}",
                text
            )
        )
    )


def replace_docx(doc, repl):

    safe_map = {
        f"{{{{{k}}}}}": str(v)
        for k, v in repl.items()
    }

    def replace_in_paragraph(paragraph):

        full_text = "".join(
            run.text for run in paragraph.runs
        )

        for k, v in safe_map.items():
            full_text = full_text.replace(k, v)

        if paragraph.runs:

            paragraph.runs[0].text = full_text

            for r in paragraph.runs[1:]:
                r.text = ""

    for p in doc.paragraphs:
        replace_in_paragraph(p)

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)




# =========================
# UI
# =========================
st.set_page_config(page_title="Contract AI", layout="wide")
st.title("📄 Contract AI Engine")

os.makedirs("temp", exist_ok=True)

# =========================
# STEP 1 (UNCHANGED)
# =========================
st.header("1️⃣ Upload documents")
st.info(
    "📌 Upload your documents. These are the raw inputs used by the AI to:\n\n"
    "- Understand the structure of the contract\n"
    "- Identify fields that need to be replaced with new information\n"
    "- Determine what type of information should be inserted into each field\n"
    "- Extract the relevant information to fill the contract fields"
)

col1, col2 = st.columns(2)

with col1:
    upload("Template Contract", "t", "template_contract")
    upload("Example Contract", "ec", "example_contract")

with col2:
    upload("Example Request", "er", "example_request")
    upload("Example Event", "ee", "example_event")

upload("New Request", "nr", "new_request")
upload("New Event", "ne", "new_event")


# =========================
# STEP 2 (UNCHANGED)
# =========================
st.header("2️⃣ Mapping")
st.info("🧠 The AI reads the example documents and learns which fields exist, their meaning, the related rules, and where the related information come from (request document or event description document).")
st.caption("This step does NOT change anything. It only builds understanding of the structure.")

if st.button("Run Mapping"):

    tpl = parse_text(st.session_state["template_contract"])
    fields = extract_template_fields(tpl)

    m = ai.build_mapping(
        fields,
        parse_text(st.session_state.get("example_contract") or ""),
        parse_text(st.session_state.get("example_request") or ""),
        parse_text(st.session_state.get("example_event") or "")
    )

    st.session_state["mapping"] = m

if st.session_state.get("mapping"):
    st.success("✅ Mapping completed! The AI has understood the contract structure. The mapped fields and the related details are listed in the following")
    st.json(st.session_state["mapping"])


# =========================
# STEP 3 (UNCHANGED)
# =========================
st.header("3️⃣ Extraction")
st.info("🔎 The AI extracts values from your new request and the new event description, based on what learned in the mapping step.")
st.caption("Output may contain small errors or mixed languages — this is expected at this stage.")

if st.button("Run Extraction"):

    require("mapping", "Run Mapping first")

    ex = ai.extract_fields(
        st.session_state["mapping"],
        parse_text(st.session_state.get("new_request") or ""),
        parse_text(st.session_state.get("new_event") or "")
    )

    st.session_state["extracted_fields"] = {
        "fields": {
            f["field_name"]: f.get("value")
            for f in ex.get("fields", [])
        }
    }

if st.session_state.get("extracted_fields"):
    st.success("📦 Extraction completed! These are the raw information fetched by the AI.")
    st.json(st.session_state["extracted_fields"])


# =========================
# STEP 3.5 - CLEAN + TRANSLATE FIELDS
# =========================
st.header("4️⃣ Cleaning")
st.info("🧹 The AI fixes grammar issues and translates everything into English.")
st.caption("Meaning is preserved. This step only improves clarity and consistency.")

def normalize_fields(fields: dict):

    prompt = f"""
You are a data cleaner.

TASK:
- Fix grammar mistakes
- Translate ALL values to English
- Do NOT change field names
- Keep meaning EXACTLY
- Return ONLY JSON

INPUT:
{json.dumps(fields, indent=2)}

OUTPUT:
{{
  "fields": {{
    "field_name": "clean English text"
  }}
}}
"""

    res = client.chat.completions.create(
        model="gpt-5",
        messages=[
            {"role": "system", "content": "Return ONLY valid JSON."},
            {"role": "user", "content": prompt}
        ]
    )

    return ai.safe_json_loads(res.choices[0].message.content)


if st.button("Run Cleaning"):

    require("extracted_fields", "Run Step 3 first")

    cleaned = normalize_fields(
        st.session_state["extracted_fields"]["fields"]
    )

    st.session_state["normalized_fields"] = cleaned["fields"]


if st.session_state.get("normalized_fields"):
    st.success("🌍 Information are ready! Everything has been cleaned and translated into English.")
    st.json(st.session_state["normalized_fields"])


# =========================
# STEP 4 (UNCHANGED)
# =========================
st.header("5️⃣ Generate Contract")
st.info("📝 The AI fills the contract template using the cleaned extracted fields.")
st.caption("Only the {{FIELDS}} are replaced. The structure of the document stays unchanged.")

if st.button("Generate"):

    require("template_contract", "Upload template contract first")
    require("extracted_fields", "Run Extraction first")

    doc = Document(st.session_state["template_contract"])

    repl = st.session_state.get("normalized_fields") or st.session_state["extracted_fields"]["fields"]

    replace_docx(doc, repl)

    out = "temp/out.docx"
    doc.save(out)

    st.session_state["final_docx"] = out
    st.success("🎉 Contract successfully generated!")

if st.session_state.get("final_docx"):
    with open(st.session_state["final_docx"], "rb") as f:
        st.download_button(
            "Download contract",
            f,
            file_name="contract.docx"
        )
